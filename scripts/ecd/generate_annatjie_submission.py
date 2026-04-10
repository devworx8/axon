from __future__ import annotations

import argparse
import sys
import tempfile
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from scripts.ecd.annatjie_visual_assets import build_cover_page, generate_visual_assets
from scripts.ecd.docx_submission_utils import (
    add_visual_evidence,
    export_pdf,
    prepend_cover_page,
    set_a4_layout,
    set_default_styles,
    strip_front_matter_until_heading,
    strip_generated_cover_page,
    strip_inserted_visual_paragraphs,
)


ROOT = Path("/home/edp/Downloads/Contact_LIst_Clive")
INPUT_DOCX = ROOT / "Annatjie_Makunyane_Unit_13853_Final_Assessor_Submission.docx"
OUTPUT_DOCX = ROOT / "Annatjie_Makunyane_Unit_13853_Submission_Ready.docx"
OUTPUT_PDF = ROOT / "Annatjie_Makunyane_Unit_13853_Submission_Ready.pdf"
ASSETS_DIR = ROOT / "ecd_assets"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Annatjie's submission-ready DOCX and PDF.")
    parser.add_argument("--input-docx", default=str(INPUT_DOCX))
    parser.add_argument("--output-docx", default=str(OUTPUT_DOCX))
    parser.add_argument("--output-pdf", default=str(OUTPUT_PDF))
    return parser.parse_args()


def _lookup_field(table, label: str) -> str:
    wanted = label.strip().lower()
    for row in table.rows:
        key = row.cells[0].text.strip().lower() if row.cells else ""
        if key == wanted and len(row.cells) > 1:
            return row.cells[1].text.strip()
    return ""


def extract_submission_metadata(doc: Document) -> dict[str, str]:
    cover_table = doc.tables[0]
    lesson_table = doc.tables[7]
    return {
        "learner_name": _lookup_field(cover_table, "Learner") or "Annatjie Makunyane",
        "centre_name": _lookup_field(cover_table, "Centre / organisation") or "Young Eagles Home Care Centre",
        "theme": _lookup_field(cover_table, "Theme") or "My Body and Healthy Habits",
        "compilation_date": _lookup_field(cover_table, "Compilation date") or "07 April 2026",
        "activity_date": _lookup_field(lesson_table, "Date") or "04 April 2026",
    }


def patch_document_properties(doc: Document, learner_name: str) -> None:
    doc.core_properties.author = learner_name
    doc.core_properties.title = f"{learner_name} Unit 13853 Submission Ready"


def build_submission(input_docx: Path, output_docx: Path, output_pdf: Path) -> None:
    doc = Document(str(input_docx))
    set_a4_layout(doc)
    set_default_styles(doc)
    meta = extract_submission_metadata(doc)
    patch_document_properties(doc, meta["learner_name"])

    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        cover_path = build_cover_page(
            ASSETS_DIR,
            temp_dir,
            learner_name=meta["learner_name"],
            centre_name=meta["centre_name"],
            theme=meta["theme"],
            activity_date=meta["activity_date"],
            compilation_date=meta["compilation_date"],
        )
        image_map = generate_visual_assets(ASSETS_DIR, temp_dir)
        strip_inserted_visual_paragraphs(doc)
        strip_generated_cover_page(doc)
        strip_front_matter_until_heading(doc, "Assessment Alignment Overview", remove_first_table=True)
        prepend_cover_page(doc, cover_path)
        add_visual_evidence(doc, image_map)
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_docx))
        export_pdf(output_docx, output_pdf)


def main() -> None:
    args = parse_args()
    build_submission(Path(args.input_docx), Path(args.output_docx), Path(args.output_pdf))
    print(f"DOCX: {args.output_docx}")
    print(f"PDF: {args.output_pdf}")


if __name__ == "__main__":
    main()
