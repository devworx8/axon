from __future__ import annotations

import argparse
import sys
import tempfile
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from scripts.ecd.docx_submission_utils import (
    add_visual_evidence,
    document_text,
    export_pdf,
    iter_text_blocks,
    prepend_cover_page,
    remove_paragraph,
    remove_table,
    set_a4_layout,
    set_default_styles,
    strip_generated_cover_page,
    strip_inserted_visual_paragraphs,
)
from scripts.ecd.mildred_visual_assets import build_cover_page, generate_visual_assets


ROOT = Path("/home/edp/Desktop/Mildred_Mathebula_Submission_Prep")
INPUT_DOCX = ROOT / "final" / "Mildred_Mathebula_Unit_13853_Final_Assessor_Submission.docx"
OUTPUT_DOCX = ROOT / "final" / "Mildred_Mathebula_Unit_13853_Submission_Ready.docx"
OUTPUT_PDF = ROOT / "final" / "Mildred_Mathebula_Unit_13853_Submission_Ready.pdf"
IMAGES_DIR = ROOT / "images"


PARAGRAPH_REWRITES = {
    "This document responds to the workbook activities and provides completed portfolio evidence for assessor review.": "I completed the workbook activities and compiled this portfolio evidence pack for final assessor submission.",
    "The workbook organises the assessment into five formative activities and one summative practical task.": "I organised this submission around the five formative activities and the one summative practical task in the workbook. In each section I expanded the required points into clear academic responses grounded in my practice at Empendulo Day Care. Where evidence was required, I included observation records, support figures and visual evidence.",
    "Developmentally appropriate practice means that planning is based on how young children learn best.": "Developmentally appropriate practice means that my planning is based on how young children learn best. At Empendulo Day Care, I start with the children's age range, prior experiences, language backgrounds, interests and developmental levels. My programme is therefore not a fixed script; it is a responsive plan that uses play, guided participation, repetition, conversation and concrete materials to help children construct understanding.",
    "In my centre, planning begins with a weekly theme and a set of realistic learning outcomes.": "In my centre, planning begins with a weekly theme and a set of realistic learning outcomes. The theme I used at Empendulo Day Care is “Healthy Food, My Body and Good Choices”. From that theme I plan story time, vocabulary games, counting and sorting tasks, role play, art activities, music and movement, outdoor play and practical routines such as handwashing before snack time.",
    "Delivery of the programme follows a balanced structure.": "Delivery of my programme follows a balanced structure. I use a short whole-group session to introduce the idea, vocabulary and key experiences for the day. Children then move into active participation through small-group tasks, play stations, guided discussion or practical demonstrations. I close with reflection, revision or a song so that learning is reinforced in an enjoyable way.",
    "Planning is aligned to the South African National Curriculum Framework for young children and to CAPS expectations in the Foundation Phase.": "My planning is aligned to the South African National Curriculum Framework for young children and to CAPS expectations in the Foundation Phase. This means that I choose activities with clear developmental aims, progression and inclusion in mind. I do not rely on school-like worksheets as the main method of learning; instead, I reach curriculum goals through age-appropriate, play-rich experiences.",
    "Children do not all learn at the same pace or in the same way.": "Children do not all learn at the same pace or in the same way. Some need more modelling, some need more time, some need enrichment and some need language support. At Empendulo Day Care, I adapt activities through smaller steps, visual prompts, peer support, simplified instructions, additional practice and extension tasks.",
    "Interests are also important.": "Children's interests are also important in my planning. When children show strong interest in pretend shopping, food preparation, market role play or counting objects, I use those interests to strengthen participation. This keeps motivation high and helps children see learning as connected to real life.",
    "An anti-bias programme also means actively responding to stereotypes.": "An anti-bias programme also means actively responding to stereotypes. If a child says that only girls cook or only boys can lead a game, I use that moment to broaden thinking. Inclusive language, fair turn-taking and shared participation help children learn respect and belonging.",
    "Observation is one of the most important professional responsibilities in ECD because it allows the educator to understand what children are doing, how they are progressing and what support they need next.": "Observation is one of my most important professional responsibilities in ECD because it allows me to understand what children are doing, how they are progressing and what support they need next. Observation is not only about noticing behaviour. It includes careful recording, analysis and planning.",
    "At Mpendulo Day Care, observation is linked to daily routines, small-group work, free play, transitions and outdoor learning.": "At Empendulo Day Care, I link observation to daily routines, small-group work, free play, transitions and outdoor learning. I record factual information first and then interpret it in the light of child development.",
    "When a child participates in a fruit-sorting activity, the educator looks beyond whether the final answer is correct.": "When a child participates in a fruit-sorting activity, I look beyond whether the final answer is correct. I notice attention span, language use, hand-eye coordination, classification skills, social confidence and persistence. I then relate these behaviours to developmental expectations so that the observation becomes meaningful.",
    "The purpose of analysis is to identify current competence and the next step in learning.": "The purpose of my analysis is to identify current competence and the next step in learning. For example, if a child can sort fruit by colour but not by size, the next task may involve comparing bigger and smaller items. If a child talks confidently in the home language but hesitates in English, I plan bilingual prompts and repeated vocabulary practice.",
    "A simple roster helps the educator ensure that every child is observed over time.": "I use a simple roster to ensure that every child is observed over time. Observation should not focus only on the loudest children or those who are struggling. Fair observation means that each learner is seen in different activities and at different times.",
    "Progress can be recorded through anecdotal notes, checklists, dated work samples, short comments on drawings, and photos taken with appropriate permission.": "I record progress through anecdotal notes, checklists, dated work samples, short comments on drawings, and photos taken with appropriate permission. These records help me describe growth over time rather than judge children on a single performance.",
    "Observation must lead to action.": "Observation must lead to action. If several children struggle to compare quantities, I repeat the concept using larger concrete materials. If children show interest in pretend shopping, the next day's plan can include a classroom shop corner with baskets, labels and price cards.",
    "Learning mediation refers to the deliberate support an educator gives so that children can participate successfully and gradually become more independent.": "Learning mediation refers to the deliberate support I give so that children can participate successfully and gradually become more independent. Mediation is grounded in theories of scaffolding and guided learning. I do not do the work for the child, but I also do not leave the child unsupported.",
    "Scaffolding means giving just enough support for the child to move to the next level.": "Scaffolding means giving just enough support for the child to move to the next level. During a fruit graph activity, I may first model how to place one picture in each box, then invite the child to continue, and finally step back once the child shows confidence.",
    "Reflection can begin in simple ways.": "Reflection can begin in simple ways. I ask: What did you do first? How did you know where to put the fruit? Which part was easy? Which part was difficult? Such questions build awareness, confidence and metacognition.",
    "Instead of praising only the finished product, the educator comments on effort, strategy and thinking.": "Instead of praising only the finished product, I comment on effort, strategy and thinking. For example: “I saw you checked carefully before you matched the card” or “You tried a different way when the first one did not work.”",
    "Children participate more purposefully when they understand the goal of an activity.": "Children participate more purposefully when they understand the goal of an activity. Before starting, I explain what the class is learning, how it will happen and why it matters in daily life.",
    "Young children should also learn simple help-seeking behaviours.": "Young children should also learn simple help-seeking behaviours. I teach them to ask a friend respectfully, look at the picture chart, listen again to instructions, use available materials, or ask me for another example.",
    "This activity focuses on how the educator manages real classroom interactions.": "This activity focuses on how I manage real classroom interactions. Effective mediation requires me to keep stories interesting, organise groups well, maintain active participation and support first and additional language development.",
    "At Mpendulo Day Care, these techniques are applied during circle time, story sessions, table activities, movement tasks and routine care experiences.": "At Empendulo Day Care, I apply these techniques during circle time, story sessions, table activities, movement tasks and routine care experiences. I remain warm, responsive and observant while still keeping the learning intention clear.",
    "The educator uses voice variation, gestures, repetition, props, prediction questions, picture discussion and role play to keep children involved.": "I use voice variation, gestures, repetition, props, prediction questions, picture discussion and role play to keep children involved. I select stories from local and culturally relevant contexts so that children recognise familiar foods, homes, family roles and community settings.",
    "Large-group teaching is kept short and lively.": "I keep large-group teaching short and lively. I alternate listening with doing, use songs to regain attention, and watch for signs of fatigue or overload. Where appropriate, I extend the activity with a quick game, matching task or movement break.",
    "The educator groups children intentionally rather than randomly.": "I group children intentionally rather than randomly. A shy child may be placed with a supportive peer, while a confident child may be paired with a learner who needs modelling. Success is created through achievable tasks, supportive language and visible encouragement.",
    "Participation is strengthened by predictable routines, action songs, short turn-taking opportunities and concrete visuals.": "Participation is strengthened by predictable routines, action songs, short turn-taking opportunities and concrete visuals. Children are not embarrassed when they are unsure; instead, I give wait time, repetition and opportunities to answer with a partner first.",
    "Home language is treated as a resource rather than a problem.": "I treat home language as a resource rather than a problem. I can introduce a new English word while also acknowledging the familiar word in Sepedi, isiZulu, Xitsonga or another local language. Meaningful repetition, visual support and shared action make additional language learning natural.",
    "Reflection is essential because it helps the educator evaluate whether teaching decisions actually supported children’s learning.": "Reflection is essential because it helps me evaluate whether my teaching decisions actually supported children's learning. Reflection should not happen only at the end of the term. In effective practice, it takes place daily and informs immediate improvement.",
    "A reflective practitioner asks what worked, what did not, why children responded in particular ways, and what should be changed next time.": "As a reflective practitioner, I ask what worked, what did not, why children responded in particular ways, and what should be changed next time. Reflection therefore links theory to practice and makes my work increasingly intentional.",
    "Analysis should move beyond opinion.": "My analysis moves beyond opinion. I identify the evidence: which children engaged, which materials helped, where the timing failed, and how the grouping affected participation.",
    "The educator’s reflection is informed by child development, scaffolding theory, inclusive education principles and knowledge of play-based learning.": "My reflection is informed by child development, scaffolding theory, inclusive education principles and knowledge of play-based learning. This prevents random decision making.",
    "For example, if a teacher shortens whole-group talk because learners respond better to concrete action, that change shows understanding of developmentally appropriate practice.": "For example, when I shorten whole-group talk because learners respond better to concrete action, that change shows my understanding of developmentally appropriate practice. When I add peer support, that reflects my knowledge of social learning and scaffolding.",
    "The workbook requires the learner to deliver a developmentally appropriate activity in the ECD centre and to include evidence of suitability, observation, mediation techniques and evaluation.": "For the summative practical task, I delivered a developmentally appropriate activity in my ECD centre and included evidence of suitability, observation, mediation techniques and evaluation. The activity below describes the lesson I conducted at Empendulo Day Care.",
    "Introduction: Children gather for circle time.": "Introduction: Children gather for circle time. I introduce a basket of fruit and vegetables, ask children to name what they recognise and sing a short greeting song linked to the theme.",
    "Discussion and story: A short culturally relevant story about shopping for healthy food is shared.": "Discussion and story: I share a short culturally relevant story about shopping for healthy food. I pause for questions, prediction and vocabulary development.",
    "Mediation: The educator models the first example, asks guiding questions, gives prompts where needed, and encourages children to explain their thinking.": "Mediation: I model the first example, ask guiding questions, give prompts where needed, and encourage children to explain their thinking.",
    "Life skills routine: Before snack or pretend snack time, the educator leads a handwashing sequence and discusses why clean hands matter.": "Life skills routine: Before snack or pretend snack time, I lead a handwashing sequence and discuss why clean hands matter.",
    "Closure and reflection: Children describe one healthy food they learned about and one thing they did in the activity.": "Closure and reflection: Children describe one healthy food they learned about and one thing they did in the activity. I record the key observations.",
    "The activity was effective because children remained engaged, materials were meaningful and different abilities could be accommodated.": "The activity was effective because children remained engaged, materials were meaningful and different abilities could be accommodated. I observed strong participation during the practical sorting and discussion phases. Children were especially responsive when I used familiar foods and allowed them to speak in the language they felt most confident in.",
    "A suggested improvement would be to prepare one extra extension station for fast finishers, such as drawing a healthy lunchbox or making a simple class graph.": "To strengthen the activity next time, I would prepare one extra extension station for fast finishers, such as drawing a healthy lunchbox or making a simple class graph. I would also keep the introduction brief so that children move quickly into hands-on action.",
    "This portfolio has been finalised for assessor review and includes learner names and March observation records used throughout the submission.": "I confirm that this portfolio has been finalised for assessor review and includes the child identifiers, March observation records and visual evidence used throughout my submission.",
    "I confirm that this portfolio has been finalised for assessor review and includes the learner identifiers, March observation records and visual evidence used throughout my submission.": "I confirm that this portfolio has been finalised for assessor review and includes the child identifiers, March observation records and visual evidence used throughout my submission.",
    "Shorten teacher talk and move earlier to practical activity": "Shorten my whole-group talk and move earlier to practical activity",
}


PHRASE_REPLACEMENTS = [
    ("Mpendulo Day Care", "Empendulo Day Care"),
    ("At Mpendulo Day Care", "At Empendulo Day Care"),
    ("The activity below shows how Mildred Mathebula structured this submission at Empendulo Day Care.", ""),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Mildred's submission-ready DOCX and PDF.")
    parser.add_argument("--input-docx", default=str(INPUT_DOCX))
    parser.add_argument("--output-docx", default=str(OUTPUT_DOCX))
    parser.add_argument("--output-pdf", default=str(OUTPUT_PDF))
    return parser.parse_args()


def rewrite_text(text: str) -> str:
    original = str(text or "")
    stripped = original.strip()
    if not stripped:
        return original
    for prefix, replacement in PARAGRAPH_REWRITES.items():
        if stripped.startswith(prefix):
            return replacement
    updated = original
    for old, new in PHRASE_REPLACEMENTS:
        updated = updated.replace(old, new)
    return updated


def apply_text_rewrites(doc: Document) -> None:
    for paragraph in iter_text_blocks(doc):
        new_text = rewrite_text(paragraph.text)
        if new_text != paragraph.text:
            paragraph.text = new_text


def patch_front_matter(doc: Document) -> None:
    doc.paragraphs[2].text = "Mildred Mathebula\nEmpendulo Day Care\nNQF Level 5"
    table = doc.tables[0]
    table.cell(0, 1).text = "Mildred Mathebula"
    table.cell(1, 1).text = "Empendulo Day Care"
    table.cell(4, 1).text = "My completed workbook answers, observation records and supporting visual evidence from my practice at Empendulo Day Care."
    doc.core_properties.author = "Mildred Mathebula"
    doc.core_properties.title = "Mildred Mathebula Unit 13853 Submission Ready"


def _strip_plain_front_page(doc: Document) -> None:
    if not doc.paragraphs:
        return
    if doc.paragraphs[0].text.strip() != "Workbook Answers and Evidence Pack":
        return
    for paragraph in list(doc.paragraphs[:6]):
        remove_paragraph(paragraph)
    if doc.tables:
        remove_table(doc.tables[0])


def verify_no_old_centre_name(doc: Document) -> None:
    for paragraph in iter_text_blocks(doc):
        if "Mpendulo Day Care" in paragraph.text or "Mpendulo" in paragraph.text:
            raise RuntimeError(f"Old centre name still present: {paragraph.text[:120]}")


def build_submission(input_docx: Path, output_docx: Path, output_pdf: Path) -> None:
    doc = Document(str(input_docx))
    set_a4_layout(doc)
    set_default_styles(doc)
    apply_text_rewrites(doc)
    patch_front_matter(doc)
    doc_text = document_text(doc)

    with tempfile.TemporaryDirectory() as tmpdir:
        temp_dir = Path(tmpdir)
        cover_page = build_cover_page(IMAGES_DIR, temp_dir, document_text=doc_text)
        image_map = generate_visual_assets(IMAGES_DIR, temp_dir)
        strip_inserted_visual_paragraphs(doc)
        strip_generated_cover_page(doc)
        _strip_plain_front_page(doc)
        prepend_cover_page(doc, cover_page)
        add_visual_evidence(doc, image_map)
        verify_no_old_centre_name(doc)
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_docx))
        export_pdf(output_docx, output_pdf)


def main() -> None:
    args = parse_args()
    build_submission(Path(args.input_docx), Path(args.output_docx), Path(args.output_pdf))
    print(f"DOCX: {args.output_docx}")
    print(f"PDF: {args.output_pdf}")


if __name__ == "__main__":
    main()
