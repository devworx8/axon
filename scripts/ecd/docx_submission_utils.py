from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph


def set_a4_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)


def set_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    if "Title" in doc.styles:
        title = doc.styles["Title"]
        title.font.name = "Aptos Display"
        title.font.size = Pt(24)


def iter_text_blocks(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def document_text(doc: Document) -> str:
    return "\n".join(block.text for block in iter_text_blocks(doc) if block.text.strip())


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_picture_after(paragraph, image_path: Path, *, width_mm: int = 174) -> None:
    new_paragraph = _insert_paragraph_after(paragraph)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def add_visual_evidence(doc: Document, image_map: dict[str, Path]) -> None:
    for paragraph in list(doc.paragraphs):
        caption = paragraph.text.strip()
        image_path = image_map.get(caption)
        if image_path and image_path.exists():
            insert_picture_after(paragraph, image_path)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_table(table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_generated_cover_page(doc: Document) -> None:
    while doc.paragraphs:
        first = doc.paragraphs[0]
        has_drawing = bool(first._p.xpath(".//w:drawing"))
        has_page_break = bool(first._p.xpath(".//w:br[@w:type='page']"))
        if has_drawing or has_page_break or not first.text.strip():
            remove_paragraph(first)
            continue
        break


def strip_inserted_visual_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if has_drawing and not paragraph.text.strip():
            remove_paragraph(paragraph)


def strip_front_matter_until_heading(doc: Document, heading_text: str, *, remove_first_table: bool = True) -> None:
    while doc.paragraphs and doc.paragraphs[0].text.strip() != heading_text:
        remove_paragraph(doc.paragraphs[0])
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        remove_paragraph(doc.paragraphs[0])
    if remove_first_table and doc.tables:
        first_table = doc.tables[0]
        header = [cell.text.strip().lower() for cell in first_table.rows[0].cells] if first_table.rows else []
        if header and header[0] in {"field", "learner"}:
            remove_table(first_table)


def prepend_cover_page(doc: Document, cover_image_path: Path) -> None:
    if not doc.paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(cover_image_path), width=Mm(174))
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return
    first = doc.paragraphs[0]
    page_break_paragraph = first.insert_paragraph_before()
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    cover_paragraph = page_break_paragraph.insert_paragraph_before()
    cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paragraph.add_run().add_picture(str(cover_image_path), width=Mm(174))


def export_pdf(docx_path: Path, output_pdf: Path) -> None:
    out_dir = output_pdf.parent
    subprocess.run(
        [
            "/usr/bin/libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = out_dir / f"{docx_path.stem}.pdf"
    if generated != output_pdf and generated.exists():
        generated.replace(output_pdf)
    if not output_pdf.exists():
        raise RuntimeError("LibreOffice did not produce the expected PDF.")


__all__ = [
    "add_visual_evidence",
    "document_text",
    "export_pdf",
    "insert_picture_after",
    "iter_text_blocks",
    "prepend_cover_page",
    "remove_paragraph",
    "remove_table",
    "set_a4_layout",
    "set_default_styles",
    "strip_front_matter_until_heading",
    "strip_generated_cover_page",
    "strip_inserted_visual_paragraphs",
]
